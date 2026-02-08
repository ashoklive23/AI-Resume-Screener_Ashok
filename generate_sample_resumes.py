"""
Sample Resume Generator for Testing the AI Resume Screener
Creates realistic sample resumes in DOCX format for different experience levels
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

class SampleResumeGenerator:
    
    def __init__(self):
        self.sample_names = [
            "Rajesh Kumar", "Priya Sharma", "Amit Patel", "Sneha Reddy", "Vikram Singh",
            "Anjali Gupta", "Karthik Iyer", "Divya Nair", "Rohan Mehta", "Kavita Joshi"
        ]
        
        self.procurement_skills_pool = [
            "Strategic Sourcing", "Vendor Management", "Contract Negotiation",
            "RFP/RFQ Management", "Cost Reduction", "Supplier Development",
            "Category Management", "SAP MM", "Oracle Procurement", "Spend Analysis",
            "Compliance Management", "Risk Management", "Supplier Quality Assurance",
            "E-Procurement", "Procure-to-Pay", "Total Cost of Ownership Analysis",
            "Lean Procurement", "Just-in-Time Procurement", "Global Sourcing",
            "Contract Management", "Supplier Scorecard Development"
        ]
        
        self.premium_skills_pool = [
            "Advanced Excel (VLOOKUP, Pivot Tables, Macros, VBA)",
            "Power BI Dashboard Development",
            "Tableau Data Visualization",
            "Python for Data Analysis",
            "SQL Database Queries",
            "Data Analytics",
            "Process Automation (RPA)"
        ]
        
        self.companies = [
            "Tata Consultancy Services", "Infosys Limited", "Wipro Technologies",
            "HCL Technologies", "Tech Mahindra", "Larsen & Toubro",
            "Reliance Industries", "Aditya Birla Group", "Mahindra Group",
            "ITC Limited", "Hindustan Unilever", "Amazon India"
        ]
        
        self.job_titles = {
            '5A': ["Procurement Analyst", "Junior Buyer", "Sourcing Analyst"],
            '5B': ["Senior Procurement Analyst", "Senior Buyer", "Category Analyst"],
            '4A': ["Procurement Specialist", "Category Manager", "Senior Category Analyst"],
            '4B': ["Assistant Procurement Manager", "Senior Category Manager", "Procurement Lead"],
            '4C': ["Procurement Manager", "Head of Procurement", "Strategic Sourcing Manager"]
        }
    
    def generate_resume(self, experience_years, filename):
        """Generate a sample resume with specified experience"""
        
        # Convert to int to avoid float issues
        experience_years = int(experience_years)
        
        # Determine band
        if experience_years <= 1:
            band = '5A'
        elif experience_years <= 3:
            band = '5B'
        elif experience_years <= 7:
            band = '4A'
        elif experience_years <= 11:
            band = '4B'
        else:
            band = '4C'
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Name
        name = random.choice(self.sample_names)
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact Info
        email = f"{name.lower().replace(' ', '.')}@email.com"
        phone = f"+91 {random.randint(70000, 99999)} {random.randint(10000, 99999)}"
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(f"{email} | {phone}")
        contact_run.font.size = Pt(11)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Professional Summary
        self._add_heading(doc, "PROFESSIONAL SUMMARY")
        summary_text = f"Results-driven procurement professional with {experience_years}+ years of experience in strategic sourcing, vendor management, and cost optimization. Proven track record in managing supplier relationships, negotiating contracts, and driving procurement excellence across diverse categories."
        doc.add_paragraph(summary_text)
        
        # Core Competencies
        self._add_heading(doc, "CORE COMPETENCIES")
        num_skills = min(8 + (experience_years // 2), len(self.procurement_skills_pool))
        skills = random.sample(self.procurement_skills_pool, num_skills)
        
        # Add premium skills for senior roles
        if experience_years >= 2:
            num_premium = min(1 + (experience_years // 4), 3)
            premium_skills = random.sample(self.premium_skills_pool, num_premium)
            skills.extend(premium_skills)
        
        skills_text = " • ".join(skills)
        doc.add_paragraph(skills_text)
        
        # Professional Experience
        self._add_heading(doc, "PROFESSIONAL EXPERIENCE")
        
        # Generate work history
        current_year = 2026
        remaining_years = experience_years
        job_count = 0
        
        while remaining_years > 0 and job_count < 4:
            years_in_role = min(random.randint(2, 4), remaining_years)
            if years_in_role == 0:
                years_in_role = 1
            
            start_year = current_year - years_in_role
            end_year = current_year
            
            # Determine appropriate job title
            if remaining_years >= 12:
                title = random.choice(self.job_titles['4C'])
            elif remaining_years >= 8:
                title = random.choice(self.job_titles['4B'])
            elif remaining_years >= 4:
                title = random.choice(self.job_titles['4A'])
            elif remaining_years >= 2:
                title = random.choice(self.job_titles['5B'])
            else:
                title = random.choice(self.job_titles['5A'])
            
            company = random.choice(self.companies)
            
            # Job header
            job_para = doc.add_paragraph()
            job_run = job_para.add_run(f"{title} | {company}")
            job_run.font.bold = True
            job_run.font.size = Pt(11)
            
            # Date range
            date_text = f"Jan {start_year} - " + ("Present" if job_count == 0 else f"Dec {end_year}")
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(date_text)
            date_run.font.italic = True
            date_run.font.size = Pt(10)
            
            # Responsibilities
            responsibilities = self._generate_responsibilities(experience_years - remaining_years + years_in_role)
            for resp in responsibilities:
                doc.add_paragraph(resp, style='List Bullet')
            
            current_year = start_year
            remaining_years -= years_in_role
            job_count += 1
            
            doc.add_paragraph()  # Spacing
        
        # Education
        self._add_heading(doc, "EDUCATION")
        doc.add_paragraph("MBA in Supply Chain Management | Indian Institute of Management | 2015")
        doc.add_paragraph("B.Tech in Mechanical Engineering | National Institute of Technology | 2013")
        
        # Certifications (for experienced candidates)
        if experience_years >= 4:
            self._add_heading(doc, "CERTIFICATIONS")
            certs = [
                "Certified Professional in Supply Management (CPSM)",
                "Certified Supply Chain Professional (CSCP)",
                "Six Sigma Green Belt"
            ]
            for cert in certs[:min(2, 1 + experience_years // 5)]:
                doc.add_paragraph(cert, style='List Bullet')
        
        # Save document
        doc.save(filename)
        print(f"✅ Generated resume: {filename} (Band: {band}, Experience: {experience_years} years)")
    
    def _add_heading(self, doc, text):
        """Add a formatted heading"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    def _generate_responsibilities(self, experience_level):
        """Generate realistic job responsibilities based on experience level"""
        base_responsibilities = [
            "Managed end-to-end procurement process for assigned categories worth $5M+ annually",
            "Negotiated contracts with suppliers achieving 15-20% cost savings",
            "Conducted RFP/RFQ processes and supplier evaluations",
            "Developed and maintained strategic supplier relationships"
        ]
        
        intermediate_responsibilities = [
            "Led cross-functional teams for strategic sourcing initiatives",
            "Implemented supplier scorecard system improving quality metrics by 25%",
            "Managed supplier risk assessment and mitigation strategies",
            "Drove process improvements using Lean Six Sigma methodologies"
        ]
        
        advanced_responsibilities = [
            "Established global sourcing strategy resulting in 30% cost reduction",
            "Led procurement transformation initiatives including e-procurement implementation",
            "Mentored team of 5-8 procurement analysts and buyers",
            "Presented procurement insights to C-level executives using Power BI dashboards"
        ]
        
        responsibilities = base_responsibilities[:3]
        
        if experience_level >= 3:
            responsibilities.extend(random.sample(intermediate_responsibilities, 2))
        
        if experience_level >= 8:
            responsibilities.extend(random.sample(advanced_responsibilities, 2))
        
        return responsibilities

def main():
    """Generate sample resumes for all bands"""
    generator = SampleResumeGenerator()
    
    print("🚀 Generating Sample Resumes for Testing...\n")
    
    # Generate resumes for each band
    samples = [
        (0.5, "sample_resume_band_5A_analyst.docx"),
        (2.5, "sample_resume_band_5B_senior_analyst.docx"),
        (5, "sample_resume_band_4A_management_trainee.docx"),
        (9, "sample_resume_band_4B_assistant_manager.docx"),
        (14, "sample_resume_band_4C_manager.docx")
    ]
    
    for years, filename in samples:
        generator.generate_resume(years, filename)
    
    print("\n✨ All sample resumes generated successfully!")
    print("📁 You can now upload these resumes to test the AI Resume Screener")

if __name__ == "__main__":
    main()
